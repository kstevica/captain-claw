"""REST handlers for the file browser."""

from __future__ import annotations

import mimetypes
from pathlib import Path
from typing import TYPE_CHECKING, Any

from aiohttp import web

if TYPE_CHECKING:
    from captain_claw.web_server import WebServer


# Extensions considered safe to display as text in the browser.
_TEXT_EXTENSIONS: set[str] = {
    ".txt", ".md", ".markdown", ".rst", ".log", ".csv", ".tsv",
    ".json", ".jsonl", ".yaml", ".yml", ".toml", ".ini", ".cfg", ".conf",
    ".py", ".js", ".ts", ".jsx", ".tsx", ".html", ".htm", ".css", ".scss",
    ".xml", ".svg", ".sql", ".sh", ".bash", ".zsh", ".fish",
    ".c", ".cpp", ".h", ".hpp", ".java", ".go", ".rs", ".rb", ".php",
    ".r", ".m", ".swift", ".kt", ".lua", ".pl", ".pm",
    ".env", ".gitignore", ".dockerignore", ".editorconfig",
    ".makefile", ".cmake", ".gradle",
}

# Filenames without extensions that are text.
_TEXT_FILENAMES: set[str] = {
    "makefile", "dockerfile", "readme", "license", "changelog",
    "authors", "contributing", "todo", "notes",
}

# Maximum size (bytes) for inline text preview.
_MAX_TEXT_PREVIEW_BYTES = 2 * 1024 * 1024  # 2 MB


def _is_text_file(path: Path) -> bool:
    """Determine if a file is likely text based on extension or name."""
    suffix = path.suffix.lower()
    if suffix in _TEXT_EXTENSIONS:
        return True
    if not suffix and path.name.lower() in _TEXT_FILENAMES:
        return True
    return False


def _enrich(logical: str, physical: str, source: str) -> dict[str, Any]:
    """Build a file metadata dict from a logical/physical path pair."""
    p = Path(physical)
    exists = p.is_file()
    size = 0
    modified: float = 0
    mime_type = ""

    if exists:
        try:
            stat = p.stat()
            size = stat.st_size
            modified = stat.st_mtime
        except OSError:
            pass
        mime_type = mimetypes.guess_type(str(p))[0] or ""

    return {
        "logical": logical,
        "physical": physical,
        "filename": p.name,
        "extension": p.suffix.lower(),
        "exists": exists,
        "size": size,
        "modified": modified,
        "mime_type": mime_type or "application/octet-stream",
        "is_text": _is_text_file(p),
        "source": source,
    }


def _is_allowed_path(physical: str) -> bool:
    """Return True if the physical path is under workspace saved/ or output/."""
    try:
        from captain_claw.config import get_config
        cfg = get_config()
        workspace = cfg.resolved_workspace_path()
        resolved = Path(physical).resolve()
        saved_root = (workspace / "saved").resolve()
        output_root = (workspace / "output").resolve()
        try:
            resolved.relative_to(saved_root)
            return True
        except ValueError:
            pass
        try:
            resolved.relative_to(output_root)
            return True
        except ValueError:
            pass
    except Exception:
        pass
    return False


async def _collect_files(server: WebServer) -> list[dict[str, Any]]:
    """Merge files from in-memory registries and persisted SQLite table."""
    seen: dict[str, dict[str, Any]] = {}  # physical_path → entry

    # 1. In-memory registries (current session).
    registries: list[tuple[str, Any]] = []

    if server.agent and getattr(server.agent, "_file_registry", None):
        registries.append(("agent", server.agent._file_registry))

    if getattr(server, "_orchestrator", None) and getattr(
        server._orchestrator, "_file_registry", None
    ):
        registries.append(("orchestrator", server._orchestrator._file_registry))

    for source, registry in registries:
        for entry in registry.list_files():
            logical = entry.get("logical", "")
            physical = entry.get("physical", "")
            if not physical or physical in seen:
                continue
            seen[physical] = _enrich(logical, physical, source)

    # 2. Persisted entries from SQLite (survives restarts).
    try:
        from captain_claw.session import get_session_manager
        sm = get_session_manager()
        persisted = await sm.list_registered_files(limit=1000)
        for row in persisted:
            physical = row.get("physical_path", "")
            if not physical or physical in seen:
                continue
            logical = row.get("logical_path", "")
            source = row.get("source", "persisted")
            seen[physical] = _enrich(logical, physical, source)
    except Exception:
        pass  # Best-effort; in-memory data is still returned.

    return sorted(seen.values(), key=lambda f: f["logical"].lower())


# ------------------------------------------------------------------
# Handlers
# ------------------------------------------------------------------


async def list_files(server: WebServer, request: web.Request) -> web.Response:
    """GET /api/files — list all registered files with metadata."""
    files = await _collect_files(server)
    return web.json_response(files)


async def list_session_files(server: WebServer, request: web.Request) -> web.Response:
    """GET /api/files/session/{session_id} — list files for a specific session.

    Merges two sources:
    1. Persisted file_registry rows for this session.
    2. Filesystem scan of workspace/saved/<category>/<session_id>/ folders
       and workspace/output/<session_id>/.
    """
    session_id = request.match_info.get("session_id", "").strip()
    if not session_id:
        return web.json_response({"error": "Missing session_id"}, status=400)

    seen: dict[str, dict[str, Any]] = {}  # physical → entry

    # 1. Registry entries for this session.
    try:
        from captain_claw.session import get_session_manager
        sm = get_session_manager()
        await sm._ensure_db()
        async with sm._db.execute(
            f"""
            SELECT {sm._FILE_REG_COLS}
            FROM file_registry
            WHERE session_id = ?
            ORDER BY registered_at DESC
            """,
            (session_id,),
        ) as cursor:
            rows = await cursor.fetchall()
        cols = [c.strip() for c in sm._FILE_REG_COLS.split(",")]
        for row in rows:
            d = dict(zip(cols, row))
            physical = d.get("physical_path", "")
            logical = d.get("logical_path", "")
            source = d.get("source", "persisted")
            if physical and physical not in seen:
                seen[physical] = _enrich(logical, physical, source)
    except Exception:
        pass

    # 2. Filesystem scan — pick up files that may not be in the registry.
    try:
        from captain_claw.config import get_config
        cfg = get_config()
        workspace = cfg.resolved_workspace_path()
        saved_root = (workspace / "saved").resolve()
        output_root = (workspace / "output").resolve()

        # Normalise session_id the same way write.py does.
        safe_id = "".join(
            c if c.isalnum() or c in "._-" else "-" for c in session_id
        ).strip("-") or "default"

        # Scan workspace/saved/<category>/<session_id>/
        _CATEGORIES = (
            "downloads", "media", "output", "scripts",
            "showcase", "skills", "tmp", "tools",
        )
        for cat in _CATEGORIES:
            cat_dir = saved_root / cat / safe_id
            if not cat_dir.is_dir():
                continue
            for child in cat_dir.rglob("*"):
                if not child.is_file():
                    continue
                phys = str(child.resolve())
                if phys in seen:
                    continue
                # Build a logical path like "scripts/my_file.py"
                rel = child.relative_to(cat_dir)
                logical = f"{cat}/{rel}"
                seen[phys] = _enrich(logical, phys, "filesystem")

        # Scan workspace/output/<session_id>/
        out_dir = output_root / safe_id
        if out_dir.is_dir():
            for child in out_dir.rglob("*"):
                if not child.is_file():
                    continue
                phys = str(child.resolve())
                if phys in seen:
                    continue
                rel = child.relative_to(out_dir)
                logical = f"output/{rel}"
                seen[phys] = _enrich(logical, phys, "filesystem")
    except Exception:
        pass

    files = sorted(seen.values(), key=lambda f: f["logical"].lower())
    return web.json_response(files)


async def get_file_content(server: WebServer, request: web.Request) -> web.Response:
    """GET /api/files/content?path=<physical_path> — read text file content."""
    physical = request.query.get("path", "").strip()
    if not physical:
        return web.json_response({"error": "Missing 'path' parameter"}, status=400)

    # Security: verify path is in registry OR under workspace saved/output dirs.
    if not _is_allowed_path(physical):
        known_physicals = {f["physical"] for f in await _collect_files(server)}
        if physical not in known_physicals:
            return web.json_response({"error": "File not in registry"}, status=403)

    p = Path(physical)
    if not p.is_file():
        return web.json_response({"error": "File not found on disk"}, status=404)

    try:
        stat = p.stat()
        if stat.st_size > _MAX_TEXT_PREVIEW_BYTES:
            return web.json_response(
                {"error": "File too large for text preview (>2 MB)"},
                status=413,
            )
    except OSError as exc:
        return web.json_response({"error": str(exc)}, status=500)

    try:
        content = p.read_text(encoding="utf-8", errors="replace")
    except Exception as exc:
        return web.json_response({"error": f"Read error: {exc}"}, status=500)

    mime_type = mimetypes.guess_type(str(p))[0] or "text/plain"
    return web.json_response({
        "path": physical,
        "filename": p.name,
        "content": content,
        "size": stat.st_size,
        "mime_type": mime_type,
    })


async def download_file(server: WebServer, request: web.Request) -> web.Response:
    """GET /api/files/download?path=<physical_path> — download any registered file."""
    physical = request.query.get("path", "").strip()
    if not physical:
        return web.json_response({"error": "Missing 'path' parameter"}, status=400)

    # Security: verify path is in registry OR under workspace saved/output dirs.
    if not _is_allowed_path(physical):
        known_physicals = {f["physical"] for f in await _collect_files(server)}
        if physical not in known_physicals:
            return web.json_response({"error": "File not in registry"}, status=403)

    p = Path(physical)
    if not p.is_file():
        return web.json_response({"error": "File not found on disk"}, status=404)

    return web.FileResponse(
        p,
        headers={
            "Content-Disposition": f'attachment; filename="{p.name}"',
        },
    )


# Allowed media extensions for the /api/media endpoint.
_MEDIA_EXTENSIONS: set[str] = {
    ".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".svg",
    ".mp3", ".wav", ".ogg", ".flac", ".m4a", ".aac",
    ".mp4", ".webm", ".mov",
}


async def serve_media(server: WebServer, request: web.Request) -> web.Response:
    """GET /api/media?path=<absolute_path> — serve media files from saved/ directory.

    Unlike /api/files/download this does NOT require file registry lookup.
    Security is enforced by requiring the resolved path to be under the
    workspace ``saved/`` or ``output/`` directories with an allowed media
    extension.
    """
    raw_path = request.query.get("path", "").strip()
    if not raw_path:
        return web.json_response({"error": "Missing 'path' parameter"}, status=400)

    # Determine workspace root from config (workspace.path, e.g. ./workspace).
    from captain_claw.config import get_config
    cfg = get_config()
    workspace = cfg.resolved_workspace_path()
    saved_root = (workspace / "saved").resolve()
    output_root = (workspace / "output").resolve()

    # Support both absolute and relative paths (relative to workspace root).
    raw = Path(raw_path)
    if raw.is_absolute():
        p = raw.resolve()
    else:
        p = (workspace / raw).resolve()

    # Security: path must be under saved/ or output/.
    try:
        p.relative_to(saved_root)
    except ValueError:
        try:
            p.relative_to(output_root)
        except ValueError:
            return web.json_response(
                {"error": "Path not under workspace saved/ or output/"},
                status=403,
            )

    # Extension check.
    if p.suffix.lower() not in _MEDIA_EXTENSIONS:
        return web.json_response(
            {"error": f"Extension '{p.suffix}' not allowed for media serving"},
            status=403,
        )

    if not p.is_file():
        return web.json_response({"error": "File not found on disk"}, status=404)

    mime = mimetypes.guess_type(str(p))[0] or "application/octet-stream"
    return web.FileResponse(
        p,
        headers={
            "Content-Type": mime,
            "Cache-Control": "private, max-age=3600",
        },
    )


# ── Markdown export (PDF / DOCX) ──────────────────────────────────────────

_EXPORT_CSS = """
body { font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
       line-height: 1.6; color: #1a1a1a; max-width: 780px; margin: 0 auto; padding: 36px 28px; }
h1 { font-size: 22px; border-bottom: 2px solid #ddd; padding-bottom: 6px; margin-top: 28px; }
h2 { font-size: 18px; border-bottom: 1px solid #eee; padding-bottom: 4px; margin-top: 24px; }
h3 { font-size: 15px; margin-top: 20px; }
p { margin: 6px 0; }
ul, ol { padding-left: 22px; }
li { margin: 3px 0; }
code { background: #f4f4f4; padding: 2px 4px; border-radius: 3px; font-size: 12px; }
pre { background: #f4f4f4; padding: 10px; border-radius: 5px; overflow-x: auto; }
pre code { background: none; padding: 0; }
table { border-collapse: collapse; width: 100%; margin: 10px 0; }
th, td { border: 1px solid #ddd; padding: 6px 10px; text-align: left; }
th { background: #f4f4f4; font-weight: 600; }
hr { border: none; border-top: 1px solid #ddd; margin: 16px 0; }
blockquote { border-left: 3px solid #ddd; margin: 10px 0; padding: 4px 14px; color: #555; }
"""


def _md_to_html(markdown_text: str) -> str:
    """Convert markdown to styled HTML document string."""
    from markdown_it import MarkdownIt

    md = MarkdownIt("commonmark", {"typographer": True}).enable("table")
    body = md.render(markdown_text)
    return (
        "<!DOCTYPE html><html><head><meta charset='utf-8'>"
        f"<style>{_EXPORT_CSS}</style></head><body>{body}</body></html>"
    )


async def export_md(server: WebServer, request: web.Request) -> web.Response:
    """POST /api/files/export — export markdown content as PDF or DOCX.

    Body JSON: { "markdown": "<raw md text>", "format": "pdf"|"docx", "filename": "opt" }
    """
    try:
        data = await request.json()
    except Exception:
        return web.json_response({"error": "Invalid JSON body"}, status=400)

    md_text = str(data.get("markdown", "")).strip()
    fmt = str(data.get("format", "pdf")).strip().lower()
    filename_base = str(data.get("filename", "document")).strip()
    if filename_base.lower().endswith(".md"):
        filename_base = filename_base[:-3]

    if not md_text:
        return web.json_response({"error": "Empty markdown content"}, status=400)
    if fmt not in ("pdf", "docx"):
        return web.json_response({"error": "Format must be 'pdf' or 'docx'"}, status=400)

    html = _md_to_html(md_text)

    if fmt == "pdf":
        import weasyprint
        pdf_bytes = weasyprint.HTML(string=html).write_pdf()
        return web.Response(
            body=pdf_bytes,
            content_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename_base}.pdf"'},
        )

    # DOCX export via python-docx
    import io
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re as _re

    doc = DocxDocument()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()

        # ── Table detection ──
        # A table starts with a pipe-delimited row followed by a separator row
        if "|" in stripped and i + 1 < len(lines):
            next_stripped = lines[i + 1].strip()
            if _re.match(r"^\|?[\s:]*-{2,}[\s:]*(\|[\s:]*-{2,}[\s:]*)+\|?\s*$", next_stripped):
                # Collect all table lines
                table_lines = [stripped]
                i += 1
                while i < len(lines) and "|" in lines[i].strip():
                    table_lines.append(lines[i].strip())
                    i += 1
                _add_md_table(doc, table_lines)
                continue

        if not stripped:
            doc.add_paragraph("")
            i += 1
            continue

        # Headings
        hm = _re.match(r"^(#{1,6})\s+(.*)", stripped)
        if hm:
            level = min(len(hm.group(1)), 4)
            doc.add_heading(_strip_md_inline(hm.group(2)), level=level)
            i += 1
            continue

        # Horizontal rules
        if _re.match(r"^[-*_]{3,}\s*$", stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run("─" * 60)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)
            i += 1
            continue

        # Unordered list
        ulm = _re.match(r"^[-*+]\s+(.*)", stripped)
        if ulm:
            doc.add_paragraph(_strip_md_inline(ulm.group(1)), style="List Bullet")
            i += 1
            continue

        # Ordered list
        olm = _re.match(r"^\d+[.)]\s+(.*)", stripped)
        if olm:
            doc.add_paragraph(_strip_md_inline(olm.group(1)), style="List Number")
            i += 1
            continue

        # Regular paragraph with inline formatting
        p = doc.add_paragraph()
        _add_formatted_runs(p, stripped)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return web.Response(
        body=buf.read(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename_base}.docx"'},
    )


def _parse_table_row(line: str) -> list[str]:
    """Parse a markdown table row into cell values."""
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [cell.strip() for cell in line.split("|")]


def _add_md_table(doc, table_lines: list[str]) -> None:
    """Add a markdown table to the DOCX document."""
    import re as _re
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml.ns import qn

    # Separate header, separator, and data rows
    header_cells = _parse_table_row(table_lines[0])
    # Skip separator line (index 1)
    data_rows = []
    for tl in table_lines[2:]:
        if _re.match(r"^\|?[\s:]*-{2,}[\s:]*(\|[\s:]*-{2,}[\s:]*)*\|?\s*$", tl):
            continue
        data_rows.append(_parse_table_row(tl))

    num_cols = len(header_cells)
    num_rows = 1 + len(data_rows)

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = "Table Grid"

    # Header row
    for ci, cell_text in enumerate(header_cells):
        if ci >= num_cols:
            break
        cell = table.rows[0].cells[ci]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(_strip_md_inline(cell_text))
        run.bold = True
        run.font.size = Pt(10)
        # Light grey background for header
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): "F2F2F2",
        })
        shading.append(shd)

    # Data rows
    for ri, row_cells in enumerate(data_rows):
        for ci, cell_text in enumerate(row_cells):
            if ci >= num_cols:
                break
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            _add_formatted_runs(p, cell_text)
            for run in p.runs:
                run.font.size = Pt(10)

    doc.add_paragraph("")  # spacing after table


def _strip_md_inline(text: str) -> str:
    """Remove markdown inline formatting for plain text contexts."""
    import re as _re
    text = _re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = _re.sub(r"__(.+?)__", r"\1", text)
    text = _re.sub(r"\*(.+?)\*", r"\1", text)
    text = _re.sub(r"_(.+?)_", r"\1", text)
    text = _re.sub(r"`(.+?)`", r"\1", text)
    text = _re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    return text


def _add_formatted_runs(paragraph, text: str) -> None:
    """Add runs with bold/italic formatting to a paragraph."""
    import re as _re
    from docx.shared import Pt

    parts = _re.split(r"(\*\*[^*]+\*\*|__[^_]+__|`[^`]+`|\*[^*]+\*|_[^_]+_)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("__") and part.endswith("__"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("_") and part.endswith("_"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)
